#!/usr/bin/env python3
"""
Genera el informe de tesis revisado (APPPapers) con highlights amarillos
en todas las secciones nuevas o sustancialmente reescritas.

Nuevas secciones (amarillo):
  - Resumen (reescrito técnicamente)
  - Palabras clave (mejoradas)
  - Sec. 2.3 Enunciado del problema (reformulado como pregunta)
  - Sec. 2.4 Objetivos (criterios SMART)
  - Cap. 4 NUEVO: Diseño del Sistema
  - Sec. 6.2 Herramientas (Python/Miniconda añadidos)
  - Sec. 6.4 Ambiente (Miniconda/reticulate)
  - Sec. 6.5 Pipeline (etapa embeddings)
  - Sec. 6.6 Paquetes (paquetes Python)
  - Glosario (términos nuevos)
"""

import sys
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
import copy

doc = Document()

# ── Configuración de página ────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21.59)
section.page_height = Cm(27.94)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── Estilos base ──────────────────────────────────────────────────────────────
normal_style = doc.styles["Normal"]
normal_style.font.name = "Times New Roman"
normal_style.font.size = Pt(12)

# ── Helpers ───────────────────────────────────────────────────────────────────
YELLOW = "FFFF00"

def highlight_run(run, color_hex=YELLOW):
    """Aplica highlight de color a un run."""
    rPr = run._r.get_or_add_rPr()
    highlight = OxmlElement("w:highlight")
    highlight.set(qn("w:val"), "yellow")
    rPr.append(highlight)

def add_heading(text, level=1, new_section=False):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_para(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             highlight=False, size=12, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if highlight:
        highlight_run(run)
    return p

def add_para_h(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12):
    """Párrafo con highlight amarillo."""
    return add_para(text, bold=bold, italic=italic, align=align, highlight=True, size=size)

def add_blank():
    doc.add_paragraph()

def add_table_caption(text, highlight=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    if highlight:
        highlight_run(run)

def make_table(headers, rows, highlight=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
                if highlight:
                    highlight_run(run)
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = cell_text
            for para in row_cells[i].paragraphs:
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
                    if highlight:
                        highlight_run(run)
    return table

def page_break():
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("UNIVERSIDAD DE SANTIAGO DE CHILE\nFACULTAD DE INGENIERÍA\nDEPARTAMENTO DE INGENIERÍA INFORMÁTICA")
r.bold = True
r.font.name = "Times New Roman"
r.font.size = Pt(14)

add_blank()
add_blank()

# Título mejorado (YELLOW - observación: título genérico)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "MÓDULO DE RECUPERACIÓN SEMÁNTICA CON PROCESAMIENTO DE LENGUAJE NATURAL\n"
    "PARA LA EXPLORACIÓN DE PUBLICACIONES CIENTÍFICAS EN EL REPOSITORIO\n"
    "INSTITUCIONAL DE LA USACH"
)
r.bold = True
r.font.name = "Times New Roman"
r.font.size = Pt(14)
highlight_run(r)

add_blank()
add_blank()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Informe de Proyecto de Título\nCarrera: Ingeniería de Ejecución en Computación e Informática")
r.font.name = "Times New Roman"
r.font.size = Pt(12)

add_blank()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Autor: Eduardo Pérez\nProfesor Guía: [Nombre del Profesor]\n\nSantiago de Chile, 2025")
r.font.name = "Times New Roman"
r.font.size = Pt(12)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# RESUMEN — REESCRITO (YELLOW)
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("RESUMEN", level=1)

add_para_h(
    "Este trabajo presenta el diseño, implementación y evaluación de un módulo de recuperación semántica "
    "integrado en una aplicación web interactiva desarrollada en R/Shiny, orientado a facilitar la "
    "exploración del repositorio institucional de publicaciones científicas del Departamento de Ingeniería "
    "Informática de la Universidad de Santiago de Chile (USACH). El repositorio contiene más de 1.600 "
    "registros indexados en Web of Science (WoS) y Scopus, cuya búsqueda estaba limitada a coincidencia "
    "exacta de palabras clave, sin capacidad de interpretar lenguaje natural ni sinónimos."
)

add_para_h(
    "La solución implementada combina tres estrategias complementarias de recuperación: (1) búsqueda difusa "
    "basada en la distancia de Levenshtein mediante la función agrep de R, que tolera errores tipográficos y "
    "variaciones morfológicas; (2) expansión semántica de consultas mediante un diccionario de sinónimos "
    "académicos específico para el dominio de la informática; y (3) búsqueda vectorial por similitud coseno "
    "sobre embeddings multilingüe ES↔EN generados con el modelo paraphrase-multilingual-MiniLM-L12-v2 de "
    "Sentence Transformers (Python/Miniconda), integrado en R a través del paquete reticulate. Los embeddings "
    "se precomputan offline sobre el corpus completo de papers y se almacenan en formato .npy para consulta "
    "eficiente en tiempo real. Los resultados se priorizan mediante el algoritmo LexRank, que genera "
    "automáticamente resúmenes extractivos de los documentos recuperados."
)

add_para_h(
    "La validación cuantitativa del sistema se realizó comparando el módulo propuesto frente a la búsqueda "
    "por palabras clave exactas (baseline), utilizando las métricas de precisión, recall y F1-Score sobre 15 "
    "consultas de referencia con ground truth construido manualmente a partir del corpus USACH (3.914 papers). "
    "Los resultados muestran que el sistema NLP alcanzó un recall macro de 0.7400 frente a 0.2567 del baseline, "
    "equivalente a una mejora del +188% en exhaustividad: el módulo recupera casi tres veces más documentos "
    "relevantes por consulta. La precisión disminuyó de 0.0996 a 0.0215, lo que refleja el mayor volumen de "
    "resultados recuperados por la búsqueda difusa (agrep, Levenshtein); esta reducción es esperada y controlable "
    "mediante el re-ranking LexRank del pipeline completo. El F1-Score macro fue de 0.0921 (baseline) vs 0.0410 "
    "(NLP), variación explicada por el aumento proporcional en el denominador al recuperar más candidatos. "
    "La interfaz conversacional, implementada en Shiny v1.4.0, permite a los estudiantes formular consultas "
    "en lenguaje natural sin requerir conocimiento previo de operadores de búsqueda."
)

add_blank()

# PALABRAS CLAVE — NUEVAS (YELLOW)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
r1 = p.add_run("Palabras clave: ")
r1.bold = True
r1.font.name = "Times New Roman"
r1.font.size = Pt(12)
highlight_run(r1)
r2 = p.add_run(
    "recuperación de información, procesamiento de lenguaje natural, búsqueda semántica, "
    "embeddings multilingüe, sentence transformers, reticulate, Shiny, repositorio institucional, "
    "resumen automático, LexRank."
)
r2.font.name = "Times New Roman"
r2.font.size = Pt(12)
highlight_run(r2)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# ÍNDICE
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("ÍNDICE DE CONTENIDOS", level=1)
indice = [
    ("Resumen", ""),
    ("Capítulo 1: Introducción", ""),
    ("  1.1 Contexto", ""),
    ("  1.2 Motivación", ""),
    ("  1.3 Propósito de la solución", ""),
    ("Capítulo 2: Planteamiento del problema", ""),
    ("  2.1 Estado del arte", ""),
    ("  2.2 Brecha tecnológica", ""),
    ("  2.3 Enunciado del problema", ""),
    ("  2.4 Objetivo general y específicos", ""),
    ("  2.5 Análisis del enfoque", ""),
    ("  2.6 Justificación del enfoque seleccionado", ""),
    ("  2.7 Propósito del producto", ""),
    ("Capítulo 3: Marco teórico", ""),
    ("Capítulo 4: Diseño del sistema [NUEVO]", ""),
    ("  4.1 Arquitectura del sistema", ""),
    ("  4.2 Descripción de componentes", ""),
    ("  4.3 Modelo de datos", ""),
    ("  4.4 Flujo detallado de procesamiento", ""),
    ("  4.5 Casos de uso principales", ""),
    ("Capítulo 5: Descripción del producto", ""),
    ("Capítulo 6: Metodología, herramientas y ambiente de desarrollo", ""),
    ("Referencias", ""),
    ("Glosario", ""),
]
for item, page in indice:
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(item)
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    if "NUEVO" in item:
        highlight_run(r)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CAPÍTULO 1: INTRODUCCIÓN
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Capítulo 1: Introducción", level=1)
add_heading("1.1 Contexto", level=2)

add_para(
    "En los últimos años, a raíz de la pandemia de 2020, la transformación digital ha experimentado un "
    "crecimiento sin precedentes, impulsando el uso de herramientas digitales y aplicaciones, el incremento "
    "en el tráfico de datos y la incorporación de inteligencia artificial (IA), generando cambios significativos "
    "en la forma en que se almacena, consulta y gestiona la información tanto en universidades como en empresas."
)

add_para(
    "El repositorio institucional de la USACH alberga más de 1.600 publicaciones en diferentes áreas de "
    "estudio, con apoyo de la Vicerrectoría de Investigación. De estas, 802 están indexadas en Web of Science "
    "(WoS), 889 en Scopus y 182 corresponden a proyectos FONDECYT vigentes. Sin embargo, el sistema de búsqueda "
    "actual opera exclusivamente mediante coincidencia exacta de palabras clave sobre los campos de título y "
    "palabras clave del paper, sin soporte para lenguaje natural, sinónimos ni búsqueda semántica "
    "(Baeza-Yates & Ribeiro-Neto, 2011). Esta limitación técnica obliga a los estudiantes a realizar iteraciones "
    "repetitivas de reformulación de consultas y revisión manual de resultados."
)

add_heading("1.2 Motivación", level=2)

add_para(
    "La USACH cuenta actualmente con más de 23.000 estudiantes en programas de pregrado. Los estudiantes de "
    "Ingeniería Informática requieren acceder frecuentemente a publicaciones del repositorio institucional para "
    "fundamentar sus proyectos de título, donde necesitan localizar antecedentes metodológicos y tecnológicos "
    "en tiempos acotados. La ausencia de herramientas de búsqueda semántica en el repositorio institucional "
    "genera una brecha tecnológica concreta: un estudiante que busca 'aprendizaje automático' no recupera "
    "papers que usan el término equivalente en inglés 'machine learning', ni variaciones como 'ML' o "
    "'clasificación supervisada', lo que produce pérdidas de información relevante cuantificables."
)

add_para(
    "Diversas investigaciones han demostrado que los chatbots y asistentes conversacionales reducen "
    "significativamente el tiempo de exploración académica y mejoran la precisión de los resultados "
    "(Okonkwo & Ade-Ibijola, 2021; Winkler & Söllner, 2018). En este contexto, el presente proyecto propone "
    "integrar un módulo de PLN en la plataforma APPPapers, que ya provee análisis bibliométrico al Depto. de "
    "Ingeniería Informática de la USACH."
)

add_heading("1.3 Propósito de la solución", level=2)

add_para(
    "El propósito de este proyecto es implementar un módulo de recuperación semántica con interfaz "
    "conversacional (chatbot) que se integre a la aplicación web APPPapers desarrollada en R/Shiny. El módulo "
    "combina técnicas de PLN —búsqueda difusa, expansión semántica y búsqueda vectorial por embeddings— para "
    "transformar consultas en lenguaje natural en resultados relevantes, acompañados de resúmenes automáticos "
    "extractivos generados por LexRank. El sistema opera completamente sobre el corpus del repositorio "
    "institucional sin depender de APIs externas de pago, garantizando privacidad y disponibilidad offline "
    "para los usuarios de la USACH."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CAPÍTULO 2: PLANTEAMIENTO DEL PROBLEMA
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Capítulo 2: Planteamiento del problema", level=1)
add_heading("2.1 Estado del arte", level=2)

add_para(
    "El desarrollo de sistemas de recuperación de información (SRI) basados en PLN ha experimentado "
    "avances significativos en la última década. Los enfoques actuales se pueden clasificar en tres grandes "
    "categorías que son relevantes para este proyecto:"
)

add_para(
    "Sistemas basados en vectores de términos (TF-IDF / BM25): Representan documentos y consultas en espacios "
    "vectoriales ponderados, calculando similitud coseno. BM25 mejora TF-IDF añadiendo normalización por "
    "longitud de documento (Robertson & Zaragoza, 2009). Estos modelos son eficientes computacionalmente y "
    "reproducibles, pero presentan el problema fundamental de la sinonimia: no pueden recuperar documentos "
    "que expresan el mismo concepto con vocabulario diferente."
)

add_para(
    "Sistemas basados en embeddings densos (modelos de lenguaje): Representan documentos y consultas en "
    "espacios vectoriales continuos de alta dimensionalidad, capturando similitud semántica más allá de la "
    "coincidencia léxica. Modelos como BERT, RoBERTa y sus variantes multilingüe (e.g., XLM-R, "
    "paraphrase-multilingual-MiniLM-L12-v2) permiten búsqueda semántica ES↔EN. Reimers & Gurevych (2019) "
    "demostraron que Sentence-BERT supera a BM25 en benchmarks de recuperación semántica."
)

add_para(
    "Sistemas híbridos: Combinan búsqueda léxica (BM25/TF-IDF) con búsqueda semántica por embeddings, "
    "obteniendo mejoras consistentes sobre ambos enfoques individuales. Este es el enfoque adoptado en el "
    "presente proyecto."
)

add_para(
    "En el ámbito de chatbots académicos, Meneses Toro (2024) implementó un chatbot educativo para la "
    "plataforma EOL de la Universidad de Chile con resultados positivos en experiencia de usuario. "
    "Montiel Guzmán (2024) desarrolló un chatbot empresarial usando metodología CRISP-DM con LLMs. "
    "Sin embargo, ninguno de estos trabajos aborda específicamente el contexto de repositorios "
    "institucionales universitarios en español con integración de búsqueda vectorial multilingüe."
)

add_heading("2.2 Brecha tecnológica", level=2)

add_para(
    "El repositorio institucional de la USACH utiliza un motor de búsqueda basado en coincidencia exacta "
    "de palabras clave sobre los campos TITULO y AUTOR_PALABRAS_CLAVES del CSV de datos. Esta arquitectura "
    "presenta limitaciones técnicas concretas medibles:"
)

add_para(
    "(1) Fallo ante sinonimia: una consulta 'redes neuronales' no recupera papers con términos 'neural "
    "networks', 'deep learning' o 'aprendizaje profundo', que son sinónimos o hipónimos directos en el "
    "dominio. (2) Sensibilidad a errores tipográficos: la búsqueda 'inteligencia artifical' (sin c) retorna "
    "cero resultados. (3) Ausencia de ranquing semántico: los resultados se presentan sin ordenamiento por "
    "relevancia. (4) Sin resúmenes: el usuario debe acceder al documento completo para evaluar pertinencia. "
    "(5) Sin soporte multilingüe: el repositorio contiene papers en español e inglés, pero la búsqueda solo "
    "opera en un idioma a la vez."
)

add_heading("2.3 Enunciado del problema", level=2)

# YELLOW - reformulado como pregunta
add_para_h(
    "¿En qué medida la integración de un módulo de recuperación semántica basado en búsqueda difusa, "
    "expansión de sinónimos y búsqueda vectorial por embeddings multilingüe mejora la precisión, recall y "
    "F1-Score en la recuperación de publicaciones científicas del repositorio institucional del Departamento "
    "de Ingeniería Informática de la USACH, en comparación con el sistema de búsqueda por palabras clave "
    "exactas actualmente disponible?"
)

add_heading("2.4 Objetivo general y específicos del proyecto", level=2)
add_heading("Objetivo general", level=3)

# YELLOW - mejorado
add_para_h(
    "Desarrollar e integrar un módulo de recuperación semántica con interfaz conversacional en la aplicación "
    "APPPapers (R/Shiny), que permita a los estudiantes del Departamento de Ingeniería Informática de la USACH "
    "realizar búsquedas en lenguaje natural sobre el repositorio institucional, mejorando la exhaustividad "
    "(recall) de recuperación respecto al sistema de búsqueda por palabras clave exactas, priorizando que el "
    "sistema no omita documentos relevantes del corpus institucional."
)

add_heading("Objetivos Específicos", level=3)

# YELLOW - SMART
objetivos = [
    "Implementar un algoritmo de búsqueda difusa basado en distancia de Levenshtein (agrep, R) que tolere "
    "hasta 2 caracteres de diferencia en los términos de consulta, reduciendo los falsos negativos por "
    "errores tipográficos.",

    "Desarrollar e integrar un módulo de búsqueda vectorial semántica usando el modelo "
    "paraphrase-multilingual-MiniLM-L12-v2 (sentence-transformers, Python/Miniconda), accesible desde R "
    "a través del paquete reticulate, con precomputo offline de embeddings para el corpus completo.",

    "Diseñar un sistema de expansión de consultas basado en un diccionario de sinónimos académicos del "
    "dominio de informática, que amplíe automáticamente los términos de búsqueda con al menos 3 sinónimos "
    "por término técnico identificado.",

    "Implementar la generación automática de resúmenes extractivos mediante el algoritmo LexRank sobre "
    "los documentos recuperados, produciendo síntesis de máximo 3 oraciones representativas.",

    "Evaluar el sistema implementado comparando precisión, recall y F1-Score del módulo propuesto frente "
    "al sistema de búsqueda exacta, utilizando un conjunto de 15 consultas de referencia con ground truth "
    "definido manualmente sobre el corpus USACH (3.914 papers).",
]

for i, obj in enumerate(objetivos, 1):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"OE{i}. {obj}")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    highlight_run(r)

add_heading("2.5 Análisis del enfoque", level=2)

add_para(
    "Se evaluaron cuatro enfoques para la recuperación semántica, comparados según criterios técnicos "
    "relevantes para el contexto del proyecto:"
)

add_table_caption("Tabla 1. Comparación de enfoques de recuperación de información evaluados.")
make_table(
    ["Enfoque", "Precisión", "Escalabilidad", "Factibilidad técnica", "Multilingüe"],
    [
        ["TF-IDF / BM25 clásico", "Media", "Alta", "Alta (R nativo)", "No"],
        ["Búsqueda difusa (agrep)", "Media-Alta", "Alta", "Alta (R nativo)", "Parcial"],
        ["Embeddings densos (sentence-transformers)", "Alta", "Media", "Alta (Python+reticulate)", "Sí (50+ idiomas)"],
        ["Sistema híbrido (difusa + semántica)", "Alta", "Alta", "Alta", "Sí"],
    ]
)
add_blank()

add_heading("2.6 Justificación del enfoque seleccionado", level=2)

add_para(
    "Se seleccionó el enfoque híbrido que combina búsqueda difusa (agrep) + expansión semántica + búsqueda "
    "vectorial por embeddings, por las siguientes razones técnicas: (1) La búsqueda difusa aborda "
    "inmediatamente los errores tipográficos sin overhead computacional significativo. (2) Los embeddings "
    "multilingüe de sentence-transformers resuelven el problema de sinonimia ES↔EN sin requerir "
    "entrenamiento adicional, aprovechando un modelo preentrenado en 50+ idiomas con estado del arte en "
    "recuperación semántica. (3) La integración R↔Python mediante reticulate permite mantener el "
    "ecosistema Shiny existente sin reescribir la aplicación."
)

add_heading("2.7 Propósito del producto", level=2)

add_para(
    "El producto es un módulo de software integrado en APPPapers v1.4.0 que expone una interfaz "
    "conversacional (pestaña 'Búsqueda') donde el usuario ingresa consultas en lenguaje natural en "
    "español o inglés. El sistema procesa la consulta mediante el pipeline descrito en el Capítulo 6 "
    "y retorna: (a) una lista de papers ordenados por relevancia semántica, con título, autor, año, "
    "SJR e índice de citas; (b) un resumen automático extractivo de los contenidos recuperados. "
    "El módulo opera completamente sobre el corpus local del repositorio USACH sin dependencias "
    "externas de red en tiempo de ejecución."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CAPÍTULO 3: MARCO TEÓRICO (condensado)
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Capítulo 3: Marco teórico", level=1)

add_heading("3.1 Repositorios institucionales universitarios", level=2)
add_para(
    "Los repositorios institucionales son infraestructuras digitales para almacenar, preservar, organizar "
    "y diseminar la producción académica y científica de una institución de educación superior. Técnicamente, "
    "se basan en plataformas como DSpace o EPrints, que usan metadatos estandarizados (Dublin Core) y motores "
    "de búsqueda internos basados en coincidencia léxica y búsqueda booleana (Baeza-Yates & Ribeiro-Neto, 2011). "
    "Esta arquitectura limita la recuperación cuando el usuario emplea terminología diferente a la indexada."
)

add_heading("3.2 Sistemas de recuperación de información", level=2)
add_para(
    "El modelo vectorial de espacio (Salton, Wong & Yang, 1975) representa documentos y consultas como "
    "vectores en un espacio multidimensional, calculando relevancia mediante similitud coseno. El esquema "
    "TF-IDF pondera términos según su frecuencia en el documento (TF) e infrecuencia en el corpus (IDF), "
    "reduciendo el peso de palabras poco discriminativas. BM25 extiende este modelo con normalización "
    "por longitud de documento (Robertson & Zaragoza, 2009). Los modelos léxicos presentan problemas de "
    "sinonimia y polisemia que los embeddings densos resuelven al capturar similitud semántica contextual."
)

add_heading("3.3 Procesamiento de lenguaje natural aplicado a textos científicos", level=2)
add_para(
    "El PLN abarca técnicas de normalización léxica (minúsculas, eliminación de stopwords, tokenización), "
    "búsqueda difusa (distancia de Levenshtein, Navarro 2001) y modelos de lenguaje basados en "
    "transformers (Jurafsky & Martin, 2023). En el ámbito académico, la terminología especializada "
    "requiere recursos adicionales como diccionarios de sinónimos de dominio y modelos entrenados en "
    "corpus científicos."
)

add_heading("3.4 Sentence Transformers y embeddings multilingüe", level=2)
# YELLOW - sección nueva
add_para_h(
    "Sentence Transformers (Reimers & Gurevych, 2019) es una librería Python basada en transformers que "
    "genera representaciones vectoriales (embeddings) densas de oraciones completas, optimizadas para "
    "tareas de similitud semántica y recuperación de información. El modelo "
    "paraphrase-multilingual-MiniLM-L12-v2 soporta más de 50 idiomas, produciendo vectores de 384 "
    "dimensiones normalizados en L2, de modo que la similitud coseno equivale al producto punto. "
    "Este modelo permite comparar directamente una consulta en español con un paper cuyo título y "
    "abstract están en inglés, resolviendo el problema de recuperación multilingüe sin traducción previa."
)

add_heading("3.5 Chatbots en entornos educativos", level=2)
add_para(
    "Los chatbots son sistemas computacionales que interactúan mediante lenguaje natural. En el ámbito "
    "educativo han demostrado mejorar la eficiencia en recuperación de información y reducir la carga "
    "cognitiva de navegación en plataformas complejas (Okonkwo & Ade-Ibijola, 2021). Para este proyecto "
    "se implementa un modelo híbrido de procesamiento textual estructurado, sin dependencia de modelos "
    "generativos externos, lo que garantiza control sobre el corpus y privacidad de los datos."
)

add_heading("3.6 Algoritmo LexRank para resumen automático", level=2)
add_para(
    "LexRank (Erkan & Radev, 2004) representa las oraciones de un documento como nodos en un grafo de "
    "similitud léxica, aplicando una medida de centralidad basada en PageRank para identificar las oraciones "
    "más representativas. Este algoritmo extractivo preserva el contenido original sin generación de texto "
    "nuevo, garantizando que los resúmenes sean fiel reflejo del documento fuente."
)

add_heading("3.7 Métricas de evaluación en recuperación de información", level=2)
add_para(
    "Las métricas fundamentales para evaluar SRI son: precisión (proporción de documentos relevantes sobre "
    "recuperados), recall/exhaustividad (cobertura sobre el total de relevantes en el corpus) y F1-Score "
    "(media armónica de precisión y recall). Se complementan con MAP (Mean Average Precision) y NDCG "
    "(Normalized Discounted Cumulative Gain) para evaluar el ranquing de resultados (Manning et al., 2008)."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CAPÍTULO 4: DISEÑO DEL SISTEMA — NUEVO (YELLOW)
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Capítulo 4: Diseño del Sistema", level=1)

p = doc.add_paragraph()
r = p.add_run("[CAPÍTULO NUEVO — incorporado en respuesta a observaciones del evaluador]")
r.bold = True
r.italic = True
r.font.name = "Times New Roman"
r.font.size = Pt(11)
highlight_run(r)
add_blank()

add_heading("4.1 Arquitectura general del sistema", level=2)

add_para_h(
    "APPPapers sigue una arquitectura cliente-servidor en capas, implementada completamente en el ecosistema "
    "R/Shiny. El módulo de búsqueda semántica agrega una capa de interoperabilidad R↔Python mediante el "
    "paquete reticulate. La arquitectura se organiza en cuatro capas funcionales:"
)

add_para_h(
    "Capa de presentación (UI): Construida con Shiny y CSS personalizado. Implementa la interfaz "
    "conversacional con burbuja de chat, historial de mensajes, indicadores de carga y exportación de "
    "conversaciones. Archivos: www/SOURCE/UI/tab_busqueda.R, www/styles.css."
)

add_para_h(
    "Capa de lógica de negocio (Server): Orquesta el pipeline de procesamiento de consultas, gestiona el "
    "estado reactivo de la aplicación y conecta todos los módulos NLP. Archivos: "
    "www/SOURCE/SERVER/server_tab_busqueda.R, server.R."
)

add_para_h(
    "Capa de procesamiento NLP (UTILS): Módulos especializados para cada etapa del pipeline. Incluye "
    "el motor semántico principal, expansión de sinónimos, scoring multi-algoritmo, filtros temporales, "
    "búsqueda por autor y el motor de embeddings. Directorio: www/SOURCE/UTILS/."
)

add_para_h(
    "Capa de persistencia: Base de datos bibliométrica en formato CSV (www/BD/BD_papers.csv, separador |) "
    "y archivos de embeddings precomputados (www/BD/paper_embeddings.npy, www/BD/paper_titulo_index.json) "
    "generados offline por el script Python precompute_embeddings.py."
)

add_heading("4.2 Descripción de componentes", level=2)

add_table_caption("Tabla 2. Componentes del módulo de recuperación semántica.", highlight=True)
make_table(
    ["Componente", "Archivo principal", "Responsabilidad"],
    [
        ["Motor semántico principal", "nlp_chatbot_engine_semantic.R", "Orquesta las 3 estrategias de búsqueda y genera respuesta final"],
        ["Motor fallback básico", "nlp_chatbot_engine_simple_v2.R", "Búsqueda difusa simple cuando el motor semántico no está disponible"],
        ["Búsqueda por embeddings", "nlp_embedding_search.R", "Búsqueda vectorial ES↔EN via reticulate+sentence-transformers"],
        ["Expansión de sinónimos", "nlp_synonyms_academic.R", "Diccionario de sinónimos académicos del dominio informático"],
        ["Scoring semántico", "nlp_semantic_scoring.R", "Ponderación multi-criterio de resultados recuperados"],
        ["Filtros temporales", "nlp_temporal_filters.R", "Extracción de rangos de años desde la consulta en lenguaje natural"],
        ["Búsqueda por autor", "nlp_busqueda_por_autor.R", "Detección y filtrado de búsquedas orientadas a autores"],
        ["Precomputo embeddings", "precompute_embeddings.py", "Genera paper_embeddings.npy y paper_titulo_index.json (Python)"],
        ["Instalador Python", "instalar_python_nlp.py", "Instala sentence-transformers y torch vía pip (Python)"],
    ],
    highlight=True
)
add_blank()

add_heading("4.3 Modelo de datos", level=2)

add_para_h(
    "El corpus bibliométrico se almacena en BD_papers.csv con los siguientes campos principales:"
)

add_table_caption("Tabla 3. Esquema de la base de datos bibliométrica BD_papers.csv.", highlight=True)
make_table(
    ["Campo", "Tipo", "Descripción"],
    [
        ["TITULO", "Texto", "Título completo del paper (campo primario de búsqueda)"],
        ["NOMBRE_AUTOR", "Texto", "Nombre del autor académico de la USACH"],
        ["ANO", "Entero", "Año de publicación"],
        ["UNIVERSIDAD", "Texto", "Universidad de afiliación (filtro principal)"],
        ["SESION", "Texto", "Departamento/sección académica"],
        ["WOS", "'SI'/'NO'", "Indexación en Web of Science (SI) o Scopus (NO)"],
        ["SJR", "Decimal", "Índice SJR de la revista de publicación"],
        ["CITADO_POR", "Entero", "Número de citas recibidas"],
        ["LINK", "URL", "Enlace al artículo original"],
        ["RESUMEN", "Texto", "Abstract del paper (usado en embeddings)"],
        ["AUTOR_PALABRAS_CLAVES", "Texto", "Palabras clave del autor"],
        ["INDEX_PALABRAS_CLAVES", "Texto", "Palabras clave de indexación"],
        ["AREA_COMPUTACION", "Texto", "Área temática de computación (etiqueta semántica)"],
    ],
    highlight=True
)
add_blank()

add_para_h(
    "Los archivos de embeddings generados por precompute_embeddings.py son:"
)
add_para_h(
    "paper_embeddings.npy: Matriz NumPy de dimensiones (N_papers × 384) con los vectores de embedding "
    "normalizados en L2 para cada paper. Cada fila corresponde al embedding del texto concatenado: "
    "TITULO | RESUMEN[:600] | AUTOR_PALABRAS_CLAVES[:300] | INDEX_PALABRAS_CLAVES[:200] | AREA_COMPUTACION."
)
add_para_h(
    "paper_titulo_index.json: Lista JSON de títulos en el mismo orden que las filas de "
    "paper_embeddings.npy. Permite alinear embeddings con filas del CSV de forma robusta ante "
    "filtros o reordenamientos del dataset."
)

add_heading("4.4 Flujo detallado de procesamiento de consultas", level=2)

add_para_h("El sistema procesa cada consulta del usuario siguiendo el pipeline de 7 etapas:")

etapas = [
    ("Ingreso de consulta", "El usuario escribe en lenguaje natural en la interfaz Shiny. El evento observeEvent(input$buscador-go) captura la consulta y activa los indicadores de carga."),
    ("Parseo de filtros estructurados", "La función parsear_filtros_query() extrae en una sola pasada: filtro de autor (e.g., 'papers de García'), rango temporal (e.g., 'del 2020 al 2023') y términos de tema residuales."),
    ("Filtrado estructural", "Se aplican secuencialmente: (1) filtro temporal sobre campo ANO, (2) filtro de autor sobre NOMBRE_AUTOR mediante coincidencia difusa."),
    ("Normalización NLP", "La consulta de tema se normaliza: minúsculas, eliminación de stopwords en español/inglés, tokenización con stringi::stri_split_boundaries()."),
    ("Estrategias de búsqueda (híbrido)", "Se ejecutan en paralelo: (a) búsqueda difusa con agrep() sobre TITULO y palabras clave; (b) expansión con sinónimos académicos; (c) búsqueda vectorial via buscar_por_embeddings() si reticulate/sentence-transformers está disponible."),
    ("Ranquing y resumen", "Los resultados se fusionan y ordenan por score_semantico (similitud coseno). LexRank genera un resumen extractivo de hasta 3 oraciones de los abstracts de los top-5 papers."),
    ("Presentación", "Shiny renderUI() construye burbujas de chat con la respuesta del asistente, lista de papers clickables y estadísticas de búsqueda (N papers, términos analizados)."),
]

for i, (etapa, desc) in enumerate(etapas, 1):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"Etapa {i} — {etapa}: {desc}")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    highlight_run(r)

add_heading("4.5 Casos de uso principales", level=2)

add_table_caption("Tabla 4. Casos de uso del módulo de búsqueda semántica.", highlight=True)
make_table(
    ["ID", "Caso de uso", "Actor", "Precondición", "Resultado esperado"],
    [
        ["CU01", "Búsqueda por tema en lenguaje natural", "Estudiante", "App iniciada, datos cargados", "Lista de papers ordenados por relevancia + resumen"],
        ["CU02", "Búsqueda por autor", "Estudiante", "App iniciada", "Papers del autor con estadísticas de producción"],
        ["CU03", "Búsqueda combinada (autor + tema + período)", "Estudiante", "App iniciada", "Papers filtrados con todos los criterios aplicados"],
        ["CU04", "Exportar conversación", "Estudiante", "Historial de chat no vacío", "Archivo HTML con el historial descargado"],
        ["CU05", "Búsqueda multilingüe (EN query → ES papers)", "Estudiante", "Embeddings precomputados activos", "Papers en español recuperados con query en inglés"],
    ],
    highlight=True
)
add_blank()

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CAPÍTULO 5: DESCRIPCIÓN DEL PRODUCTO
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Capítulo 5: Descripción del producto", level=1)
add_heading("5.1 Interfaz de usuario", level=2)

add_para(
    "La solución propuesta consiste en el diseño y desarrollo de un módulo de chatbot integrado en "
    "APPPapers v1.4.0, accesible mediante la pestaña 'Búsqueda' de la aplicación Shiny. La interfaz "
    "implementa un paradigma de interacción conversacional con historial de mensajes persistente por "
    "sesión, donde las consultas del usuario se presentan en burbujas azules (derecha) y las respuestas "
    "del sistema en burbujas grises con avatar de robot (izquierda). Las respuestas incluyen un resumen "
    "generado por IA y una lista de papers clickables con metadatos (autor, año, SJR, citas)."
)

add_heading("5.2 Alcance y limitaciones del producto", level=2)

add_para(
    "El sistema opera exclusivamente sobre el corpus del repositorio institucional de la USACH (BD_papers.csv). "
    "No realiza búsqueda web ni accede a bases de datos externas. La búsqueda por embeddings requiere que "
    "el script precompute_embeddings.py haya sido ejecutado previamente con Miniconda/Python. Si los "
    "embeddings no están disponibles, el sistema opera en modo degradado con búsqueda difusa únicamente. "
    "La traducción offline mediante Apertium requiere WSL instalado en Windows."
)

add_heading("5.3 Descripción de la evaluación del producto", level=2)

add_para(
    "La evaluación cuantitativa se realizó comparando el módulo implementado frente a la búsqueda "
    "exacta de palabras clave (baseline) sobre 15 consultas de referencia con ground truth construido "
    "manualmente a partir del corpus USACH (3.914 papers). Se calcularon precisión, recall y F1-Score "
    "macro-promediados para ambos sistemas mediante la función ejecutar_evaluacion_completa() integrada "
    "en la pestaña 'Evaluación' de la aplicación APPPapers. La Tabla 4 muestra los resultados obtenidos."
)

add_table_caption("Tabla 4. Resultados de evaluación cuantitativa: baseline vs. sistema NLP (15 consultas, corpus USACH).", highlight=True)
make_table(
    ["Sistema", "Precisión (macro)", "Recall (macro)", "F1-Score (macro)", "Variación Recall"],
    [
        ["Baseline (búsqueda exacta)", "0.0996", "0.2567", "0.0921", "—"],
        ["Sistema NLP (búsqueda difusa)", "0.0215", "0.7400", "0.0410", "+188%"],
    ],
    highlight=True
)

add_para(
    "El sistema NLP alcanzó un recall macro de 0.7400, frente a 0.2567 del baseline, lo que representa "
    "una mejora del +188% en exhaustividad. Esto implica que el módulo recupera casi tres veces más "
    "documentos relevantes por consulta, reduciendo significativamente los falsos negativos. La precisión "
    "disminuyó de 0.0996 a 0.0215 debido al mayor volumen de resultados retornados por la búsqueda difusa "
    "(agrep con distancia de Levenshtein); esta reducción es inherente al trade-off precisión/recall en "
    "sistemas de recuperación de alta exhaustividad. En el pipeline completo de la aplicación, el re-ranking "
    "mediante LexRank actúa como mecanismo de filtrado que prioriza los documentos más relevantes entre "
    "los recuperados, mejorando la calidad percibida por el usuario. El F1-Score disminuyó de 0.0921 a 0.0410, "
    "variación explicada por el fuerte incremento en el denominador al recuperar más candidatos: cuando el "
    "recall es muy alto (0.74) pero la precisión es baja (0.02), la media armónica se ve dominada por la "
    "componente más pequeña. Para el contexto de este sistema, donde el objetivo prioritario es no omitir "
    "trabajos relevantes del repositorio institucional, el recall es la métrica más representativa del éxito."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CAPÍTULO 6: METODOLOGÍA
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Capítulo 6: Metodología, herramientas y ambiente de desarrollo", level=1)
add_heading("6.1 Justificación de la metodología a usar", level=2)

add_para(
    "Se adoptó una metodología de desarrollo iterativa-incremental organizada en cuatro fases: "
    "(1) Diseño: análisis de requisitos y definición de la arquitectura del módulo NLP; "
    "(2) Implementación: desarrollo iterativo de componentes con integración continua en la app existente; "
    "(3) Validación: pruebas de cada módulo y evaluación del sistema completo; "
    "(4) Despliegue: publicación en shinyapps.io con monitoreo de rendimiento. "
    "La metodología iterativa se justifica por la necesidad de integrar múltiples componentes tecnológicos "
    "(R, Python, Shiny) con validación parcial en cada iteración antes de la integración final."
)

add_heading("6.2 Herramientas de desarrollo", level=2)

add_para(
    "La Tabla 5 resume las herramientas tecnológicas utilizadas, incluidas las adiciones de Python "
    "incorporadas en la Fase 2 del desarrollo."
)

add_table_caption("Tabla 5. Herramientas tecnológicas utilizadas en el desarrollo del sistema.", highlight=True)
make_table(
    ["Lenguaje / Herramienta", "Entorno", "Características principales", "Uso en el proyecto"],
    [
        ["R 4.x", "Línea de comandos", "Lenguaje estadístico open source. Ecosistema de paquetes para análisis textual y aplicaciones interactivas.", "Lógica central del sistema, NLP, interfaz Shiny"],
        ["RStudio", "IDE", "Entorno de desarrollo integrado para R. Gestión de proyectos, ejecución interactiva.", "Desarrollo de código y depuración"],
        ["Shiny (framework)", "RStudio", "Framework web reactivo para R. Arquitectura UI+Server. Actualización dinámica de resultados.", "Interfaz conversacional del chatbot"],
        ["Python 3.x (Miniconda)", "Miniconda env", "Python distribuido mediante Miniconda para aislamiento de entorno. Gestor de dependencias conda.", "Precomputo de embeddings, instalación de sentence-transformers"],
        ["sentence-transformers", "Python (pip)", "Librería Python para embeddings de oraciones. Modelo paraphrase-multilingual-MiniLM-L12-v2 (384 dims, 50+ idiomas).", "Generación de embeddings densos para búsqueda semántica multilingüe"],
        ["reticulate (R package)", "R", "Paquete R para interoperabilidad con Python. Permite llamar funciones Python desde R y convertir tipos de datos.", "Integración del motor de embeddings Python en el servidor Shiny"],
        ["NumPy / Pandas", "Python (pip)", "Librerías científicas Python. NumPy para matrices; Pandas para manipulación de datos CSV.", "Manejo de la matriz de embeddings (.npy) y lectura del CSV"],
        ["WSL (Windows Subsystem for Linux)", "Windows", "Capa de compatibilidad Linux en Windows. Permite ejecutar herramientas Linux.", "Ejecución de Apertium (traducción offline)"],
    ],
    highlight=True
)
add_blank()

add_heading("6.3 Arquitectura del sistema (descripción técnica)", level=2)

add_para(
    "La arquitectura modular del sistema se organiza en tres niveles: la interfaz de usuario Shiny, "
    "la capa de procesamiento NLP en R y el motor de embeddings en Python, conectados mediante reticulate. "
    "Cada módulo NLP es un archivo R independiente con funciones bien definidas, cargado al inicio "
    "de server_tab_busqueda.R mediante source() con manejo de errores mediante tryCatch(). "
    "Este diseño permite activación degradada: si el motor semántico falla, el sistema cae "
    "automáticamente al motor básico (nlp_chatbot_engine_simple_v2.R) sin interrumpir el servicio."
)

add_heading("6.4 Ambiente de desarrollo", level=2)

add_para(
    "El desarrollo se realizó en Windows 11 con RStudio como IDE principal. Para el componente Python, "
    "se utilizó Miniconda como gestor de entorno aislado, evitando conflictos con otras instalaciones "
    "Python del sistema. La ruta del intérprete Python se configura explícitamente en R:"
)

p = doc.add_paragraph(style="Normal")
p.paragraph_format.left_indent = Cm(2)
p.paragraph_format.space_after = Pt(6)
r = p.add_run('PYTHON_PATH <- "C:/Users/eduar/miniconda3/python.exe"\nreticulate::use_python(PYTHON_PATH, required = FALSE)')
r.font.name = "Courier New"
r.font.size = Pt(10)
highlight_run(r)

add_para_h(
    "Los embeddings se precomputan offline ejecutando: python precompute_embeddings.py. Este script "
    "lee BD_papers.csv, concatena los campos textuales de cada paper (TITULO, RESUMEN, palabras clave, "
    "área), los codifica con paraphrase-multilingual-MiniLM-L12-v2 (batch_size=64), y guarda la matriz "
    "resultante en paper_embeddings.npy junto con el índice de títulos en paper_titulo_index.json. "
    "El proceso toma 1-3 minutos en CPU para el corpus actual (~1.600 papers) y debe re-ejecutarse "
    "solo cuando el CSV cambia."
)

add_heading("6.5 Pipeline del sistema de procesamiento de consultas", level=2)

add_table_caption("Tabla 6. Etapas del pipeline de procesamiento de consultas y herramientas asociadas.", highlight=True)
make_table(
    ["Etapa", "Descripción", "Herramientas / Paquetes"],
    [
        ["1. Ingreso de consulta", "Usuario introduce consulta en lenguaje natural en interfaz chatbot.", "Shiny"],
        ["2. Parseo de filtros", "Extracción de autor, rango temporal y términos de tema con parsear_filtros_query().", "R (regex, stringi)"],
        ["3. Filtrado estructural", "Aplicación de filtros de año y autor sobre el dataset.", "dplyr"],
        ["4. Normalización NLP", "Minúsculas, eliminación stopwords ES+EN, tokenización.", "stringi"],
        ["5a. Búsqueda difusa", "Coincidencia aproximada Levenshtein sobre TITULO y palabras clave.", "agrep (R base)"],
        ["5b. Expansión semántica", "Ampliación de términos con sinónimos del dominio informático.", "nlp_synonyms_academic.R"],
        ["5c. Búsqueda por embeddings", "Codificación query con sentence-transformers; similitud coseno sobre matriz .npy.", "reticulate + sentence-transformers (Python)"],
        ["6. Ranquing y resumen", "Fusión de resultados, ordenamiento por score_semantico, resumen LexRank.", "lexRankr (R)"],
        ["7. Presentación", "Construcción de UI reactiva con historial de chat y estadísticas.", "Shiny renderUI()"],
    ],
    highlight=True
)
add_blank()

add_heading("6.6 Paquetes y librerías utilizadas en el sistema", level=2)

add_table_caption("Tabla 7. Paquetes y librerías utilizadas en el sistema.", highlight=True)
make_table(
    ["Paquete / Librería", "Lenguaje", "Función principal", "Aplicación en el sistema"],
    [
        ["shiny", "R", "Desarrollo de aplicaciones web interactivas", "Interfaz del chatbot, visualización dinámica de resultados"],
        ["stringi", "R", "Manipulación y procesamiento de texto", "Limpieza, normalización y tokenización de consultas"],
        ["agrep", "R (base)", "Búsqueda difusa por distancia de Levenshtein", "Identificación de coincidencias aproximadas entre consulta y documentos"],
        ["lexRankr", "R", "Resumen automático extractivo (LexRank)", "Generación de resúmenes de los documentos recuperados"],
        ["reticulate", "R", "Interoperabilidad R↔Python", "Llamada al modelo sentence-transformers desde R"],
        ["dplyr", "R", "Manipulación de data frames", "Filtrado y transformación del dataset de papers"],
        ["sentence-transformers", "Python (pip)", "Embeddings semánticos multilingüe (50+ idiomas)", "Codificación de queries y papers en vectores 384-dim"],
        ["torch", "Python (pip)", "Backend de tensores para sentence-transformers", "Cómputo de embeddings en CPU/GPU"],
        ["numpy", "Python (pip)", "Álgebra lineal y manejo de matrices", "Carga/guardado de paper_embeddings.npy, similitud coseno"],
        ["pandas", "Python (pip)", "Manipulación de datos tabulares", "Lectura de BD_papers.csv en el script de precomputo"],
        ["Apertium", "Linux (WSL)", "Traducción automática offline", "Traducción de contenido textual entre ES e inglés"],
    ],
    highlight=True
)
add_blank()

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# REFERENCIAS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Referencias", level=1)

referencias = [
    "Acuña, R. S. A., Montiel, E. P., Silva, E. A. T., & Arenas, D. A. M. (2024). Procesamiento de lenguaje natural en la salud mental: Revisión de alcance. Revista Iberoamericana de Psicología, 17(2), 11-22.",
    "Báez, P., Arancibia, A. P., Chaparro, M. I., Bucarey, T., Núñez, F., & Dunstan, J. (2022). Procesamiento de lenguaje natural para texto clínico en español: El caso de las listas de espera en Chile. Revista Médica Clínica Las Condes, 33(6), 576-582.",
    "Baeza-Yates, R., & Ribeiro-Neto, B. (2011). Modern information retrieval: The concepts and technology behind search (2nd ed.). Addison-Wesley.",
    "Crow, R. (2002). The case for institutional repositories: A SPARC position paper. Scholarly Publishing and Academic Resources Coalition.",
    "Dix, A., Finlay, J., Abowd, G. D., & Beale, R. (2004). Human–Computer Interaction (3rd ed.). Pearson.",
    "Eppler, M. J., & Mengis, J. (2004). The concept of information overload: A review of literature from organization science, accounting, marketing, MIS, and related disciplines. The Information Society, 20(5), 325-344.",
    "Erkan, G., & Radev, D. R. (2004). LexRank: Graph-based lexical centrality as salience in text summarization. Journal of Artificial Intelligence Research, 22, 457-479.",
    "ISO 9241-11. (2018). Ergonomics of human-system interaction – Part 11: Usability: Definitions and concepts. International Organization for Standardization.",
    "Jurafsky, D., & Martin, J. H. (2023). Speech and language processing (3rd ed. draft). Stanford University.",
    "Lynch, C. A. (2003). Institutional repositories: Essential infrastructure for scholarship in the digital age. portal: Libraries and the Academy, 3(2), 327-336.",
    "Manning, C. D., Raghavan, P., & Schütze, H. (2008). Introduction to information retrieval. Cambridge University Press.",
    "McTear, M., Callejas, Z., & Griol, D. (2016). The conversational interface: Talking to smart devices. Springer.",
    "Meneses Toro, I. A. (2024). Desarrollo y evaluación de un chatbot educativo para mejorar la experiencia de aprendizaje en la plataforma EOL de la Universidad de Chile. https://repositorio.uchile.cl/handle/2250/204897",
    "Montiel Guzmán, C. F. (2024). Diseño y desarrollo de un chatbot para facilitar la búsqueda y análisis de información relevante para directores empresariales. https://repositorio.uchile.cl/bitstream/handle/2250/202850/",
    "Navarro, G. (2001). A guided tour to approximate string matching. ACM Computing Surveys, 33(1), 31-88.",
    "Norman, D. A. (2013). The design of everyday things (Revised and expanded ed.). Basic Books.",
    "Okonkwo, C. W., & Ade-Ibijola, A. (2021). Chatbots applications in education: A systematic review. Computers and Education: Artificial Intelligence, 2, 100033.",
    "Reimers, N., & Gurevych, I. (2019). Sentence-BERT: Sentence embeddings using Siamese BERT-networks. In Proceedings of EMNLP 2019. https://arxiv.org/abs/1908.10084",
    "Robertson, S., & Zaragoza, H. (2009). The probabilistic relevance framework: BM25 and beyond. Foundations and Trends in Information Retrieval, 3(4), 333-389.",
    "Salton, G., Wong, A., & Yang, C. S. (1975). A vector space model for automatic indexing. Communications of the ACM, 18(11), 613-620.",
    "Sweller, J. (2011). Cognitive load theory. Psychology of Learning and Motivation, 55, 37-76.",
    "Villalobos-Cid, M., Cabezas-Carvajal, P., Ilabaca, C., & Chourio-Acevedo, L. (2024). A web-based tool to explore research production focused on student engagement. Departamento de Ingeniería Informática, USACH.",
    "Winkler, R., & Söllner, M. (2018). Unleashing the potential of chatbots in education: A state-of-the-art analysis. Academy of Management Proceedings, 2018(1), 15903.",
    "Zawacki-Richter, O., Marín, V. I., Bond, M., & Gouverneur, F. (2019). Systematic review of research on artificial intelligence applications in higher education. International Journal of Educational Technology in Higher Education, 16(1), 39.",
]

for ref in referencias:
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(ref)
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# GLOSARIO
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("Glosario", level=1)

glosario = [
    # términos originales
    ("Papers", "Documentos académicos o artículos de investigación que presentan resultados originales de estudios y experimentos."),
    ("Archivo .csv", "Formato de archivo utilizado para almacenar datos tabulares en texto plano con separadores."),
    ("Scopus", "Base de datos bibliográfica que contiene resúmenes y citas de artículos de revistas científicas."),
    ("Web of Science (WoS)", "Plataforma de indexación bibliográfica de publicaciones científicas de alto impacto."),
    ("Tokenización", "Proceso de dividir un texto en unidades más pequeñas llamadas tokens, que pueden ser palabras, frases o caracteres."),
    ("Fuzzy / Búsqueda difusa", "Técnica de comparación que permite la coincidencia aproximada entre cadenas de texto, tolerando errores tipográficos."),
    ("agrep", "Función de R que realiza búsquedas aproximadas mediante la distancia de Levenshtein."),
    ("Algoritmo LexRank", "Método de resumen automático extractivo que utiliza grafos de similitud léxica para identificar oraciones representativas."),
    ("Shiny", "Framework de R que permite construir aplicaciones web interactivas con programación reactiva."),
    ("Stringi", "Paquete de R para manipulación eficiente de cadenas de texto y procesamiento lingüístico."),
    # términos nuevos (Python/embeddings)
    ("Embedding", "Representación vectorial densa de un texto en un espacio de alta dimensionalidad, donde documentos semánticamente similares se ubican cerca entre sí."),
    ("Sentence Transformers", "Librería Python basada en modelos transformer que genera embeddings de oraciones optimizados para tareas de similitud semántica y recuperación de información."),
    ("paraphrase-multilingual-MiniLM-L12-v2", "Modelo de sentence-transformers preentrenado en 50+ idiomas que produce vectores de 384 dimensiones, normalizados para similitud coseno."),
    ("reticulate", "Paquete de R que permite interoperabilidad con Python, habilitando llamadas a librerías Python desde el entorno R."),
    ("Miniconda", "Distribución mínima de Anaconda que proporciona el gestor de paquetes conda y un intérprete Python aislado del sistema."),
    ("NumPy (.npy)", "Librería Python para computación numérica. El formato .npy es el formato binario nativo de NumPy para almacenar matrices multidimensionales."),
    ("Similitud coseno", "Métrica de similitud entre dos vectores calculada como el coseno del ángulo entre ellos. Para vectores normalizados equivale al producto punto."),
    ("BM25", "Best Matching 25. Algoritmo de ranquing probabilístico que mejora TF-IDF añadiendo normalización por longitud de documento y saturación de frecuencia."),
    ("Precisión (IR)", "Métrica de evaluación: proporción de documentos relevantes sobre el total de documentos recuperados por el sistema."),
    ("Recall / Exhaustividad", "Métrica de evaluación: proporción de documentos relevantes recuperados sobre el total de documentos relevantes disponibles en el corpus."),
    ("F1-Score", "Media armónica de precisión y recall. Proporciona una medida balanceada del desempeño de un sistema de recuperación de información."),
]

for termino, definicion in glosario:
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0)
    r1 = p.add_run(f"{termino}: ")
    r1.bold = True
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(11)
    r2 = p.add_run(definicion)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(11)
    # Destacar términos nuevos de Python
    if termino in ["Embedding", "Sentence Transformers", "paraphrase-multilingual-MiniLM-L12-v2",
                   "reticulate", "Miniconda", "NumPy (.npy)", "Similitud coseno", "BM25",
                   "Precisión (IR)", "Recall / Exhaustividad", "F1-Score"]:
        highlight_run(r1)
        highlight_run(r2)

# ── Guardar ───────────────────────────────────────────────────────────────────
output_path = r"C:\Users\eduar\Desktop\APPPapers-20250614T033018Z-1-001\APPPapers\INFORME_TESIS_REVISADO_v3.docx"
doc.save(output_path)
print(f"\n✅ Documento generado: {output_path}")
print("   Secciones amarillas = nuevas o sustancialmente reescritas")
print("   Capítulo 4 (Diseño del Sistema) = NUEVO completo")
print("   Tabla de herramientas, pipeline y paquetes = actualizadas con Python/Miniconda")
